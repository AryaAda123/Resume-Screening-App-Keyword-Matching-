# import streamlit as st
# import PyPDF2
# import docx
# from nltk.corpus import stopwords
# import re # रेगुलर एक्सप्रेशन के लिए

# # --- 1. टेक्स्ट एक्सट्रैक्शन फ़ंक्शंस ---

# def extract_text_from_pdf(pdf_file):
#     """PDF फ़ाइल से टेक्स्ट निकालता है।"""
#     text = ""
#     try:
#         pdf_reader = PyPDF2.PdfReader(pdf_file)
#         for page in pdf_reader.pages:
#             text += page.extract_text() or "" # अगर पेज खाली हो तो खाली स्ट्रिंग जोड़ें
#     except Exception as e:
#         st.error(f"PDF से टेक्स्ट निकालने में त्रुटि: {e}")
#     return text

# def extract_text_from_docx(docx_file):
#     """DOCX फ़ाइल से टेक्स्ट निकालता है।"""
#     text = ""
#     try:
#         doc = docx.Document(docx_file)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         st.error(f"DOCX से टेक्स्ट निकालने में त्रुटि: {e}")
#     return text

# def extract_text_from_txt(txt_file):
#     """TXT फ़ाइल से टेक्स्ट निकालता है।"""
#     return txt_file.read().decode("utf-8")

# # --- 2. टेक्स्ट प्रोसेसिंग और कीवर्ड मैचिंग ---

# def preprocess_text(text):
#     """टेक्स्ट को लोअरकेस करता है, गैर-अक्षर/संख्या वर्णों को हटाता है, स्टॉपवर्ड्स हटाता है।"""
#     text = text.lower()
#     # केवल अक्षर, संख्या और स्पेस रखें
#     text = re.sub(r'[^a-z0-9\s]', '', text)
#     words = text.split()
#     # स्टॉपवर्ड्स हटाना (वैकल्पिक, यदि आप चाहते हैं तो हटा सकते हैं)
#     # stop_words = set(stopwords.words('english'))
#     # words = [word for word in words if word not in stop_words]
#     return " ".join(words)

# def calculate_keyword_match_score(resume_text, job_description_keywords):
#     """
#     रिज्यूमे टेक्स्ट में जॉब डिस्क्रिप्शन के कीवर्ड्स को ढूंढता है
#     और एक मैच स्कोर की गणना करता है।
#     """
#     score = 0
#     matched_keywords = []

#     # सुनिश्चित करें कि कीवर्ड्स और रिज्यूमे टेक्स्ट दोनों संसाधित (processed) हों
#     processed_resume_text = preprocess_text(resume_text)
#     processed_jd_keywords = [preprocess_text(kw) for kw in job_description_keywords if kw.strip()] # खाली कीवर्ड हटा दें

#     # प्रत्येक आवश्यक कीवर्ड के लिए देखें कि क्या यह रिज्यूमे में मौजूद है
#     for keyword in processed_jd_keywords:
#         if keyword in processed_resume_text:
#             score += 1
#             matched_keywords.append(keyword)
#     return score, matched_keywords

# # --- Streamlit UI ---

# st.set_page_config(layout="wide")
# st.title("👨‍💻 रिज्यूमे स्क्रीनिंग ऐप (कीवर्ड मैचिंग)")
# st.markdown("---")

# st.subheader("1. जॉब आवश्यकताएं दर्ज करें")
# job_keywords_input = st.text_area(
#     "जॉब के लिए आवश्यक स्किल्स/कीवर्ड्स को कॉमा या नई लाइन से अलग करके दर्ज करें:",
#     "Python, Machine Learning, Data Science, SQL, Cloud, Communication, Problem Solving"
# )
# # कीवर्ड्स को लिस्ट में विभाजित करें और खाली स्ट्रिंग हटा दें
# required_keywords = [kw.strip() for kw in job_keywords_input.split(',') if kw.strip()]
# if not required_keywords:
#     st.warning("कृपया जॉब कीवर्ड्स दर्ज करें।")
# else:
#     st.info(f"**पहचाने गए कीवर्ड्स:** {', '.join(required_keywords)}")

# st.subheader("2. रिज्यूमे अपलोड करें")
# uploaded_resumes = st.file_uploader(
#     "PDF, DOCX या TXT फॉर्मेट में रिज्यूमे अपलोड करें",
#     type=["pdf", "docx", "txt"],
#     accept_multiple_files=True
# )

# if uploaded_resumes and required_keywords:
#     st.markdown("---")
#     st.subheader("3. स्क्रीनिंग परिणाम")

#     results = []
#     for uploaded_file in uploaded_resumes:
#         file_name = uploaded_file.name
#         file_type = file_name.split('.')[-1].lower()
#         resume_text = ""

#         with st.spinner(f"'{file_name}' से टेक्स्ट निकाल रहा है..."):
#             if file_type == "pdf":
#                 resume_text = extract_text_from_pdf(uploaded_file)
#             elif file_type == "docx":
#                 resume_text = extract_text_from_docx(uploaded_file)
#             elif file_type == "txt":
#                 resume_text = extract_text_from_txt(uploaded_file)
#             else:
#                 st.warning(f"समर्थित फ़ाइल प्रकार नहीं: {file_type} ({file_name})")
#                 continue

#         if resume_text:
#             score, matched_kws = calculate_keyword_match_score(resume_text, required_keywords)
#             results.append({
#                 "file_name": file_name,
#                 "score": score,
#                 "matched_keywords": matched_kws,
#                 "total_required_keywords": len(required_keywords)
#             })
#         else:
#             st.warning(f"'{file_name}' से टेक्स्ट नहीं निकाल सका।")

#     # परिणामों को स्कोर के आधार पर सॉर्ट करें (उच्चतम स्कोर पहले)
#     results.sort(key=lambda x: x["score"], reverse=True)

#     if results:
#         for i, res in enumerate(results):
#             st.success(f"### {i+1}. {res['file_name']}")
#             st.write(f"**मैच स्कोर:** {res['score']} / {res['total_required_keywords']}")
#             if res['matched_keywords']:
#                 st.write(f"**मैच हुए कीवर्ड्स:** {', '.join(res['matched_keywords'])}")
#             else:
#                 st.write("**कोई कीवर्ड मैच नहीं हुआ।**")
#             st.markdown("---")
#     else:
#         st.info("कोई रिज्यूमे अपलोड नहीं किया गया या संसाधित नहीं किया जा सका।")


import streamlit as st
import PyPDF2
import docx
from nltk.corpus import stopwords
import re # For regular expressions

# --- Custom CSS for a more colorful and responsive UI with animations ---
st.markdown(
    """
    <style>
    /* General container padding */
    .reportview-container .main .block-container {
        padding-top: 2rem;
        padding-right: 2rem;
        padding-left: 2rem;
        padding-bottom: 2rem;
    }

    /* Sidebar background and styling */
    .css-1d391kg { /* This class targets the sidebar */
        background-color: #e0f7fa; /* Light cyan */
        border-right: 2px solid #00BCD4; /* Cyan border */
        box-shadow: 2px 0 5px rgba(0,0,0,0.1);
    }

    /* Main Title Styling */
    h1 {
        color: #4CAF50; /* Green for main title */
        text-align: center;
        font-size: 2.8em; /* Slightly larger */
        margin-bottom: 0.6em;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1); /* Subtle shadow */
        animation: fadeIn 1s ease-out; /* Fade-in animation */
    }

    /* Subheader Styling */
    h2 {
        color: #2196F3; /* Blue for subheaders */
        font-size: 2em; /* Slightly larger */
        border-bottom: 2px solid #2196F3; /* Underline effect */
        padding-bottom: 0.3em;
        margin-top: 1.5em;
        animation: slideInLeft 0.8s ease-out; /* Slide-in animation */
    }
    h3 {
        color: #FF9800; /* Orange for resume names */
        font-size: 1.6em;
        animation: fadeIn 1.2s ease-out; /* Fade-in animation */
    }

    /* Button Styling */
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border-radius: 12px;
        border: none;
        padding: 12px 25px; /* Slightly larger padding */
        font-size: 17px; /* Slightly larger font */
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1); /* Subtle shadow */
    }
    .stButton>button:hover {
        background-color: #45a049;
        transform: translateY(-3px) scale(1.02); /* Lift and slight scale on hover */
        box-shadow: 0 6px 8px rgba(0,0,0,0.15);
    }

    /* Alert/Message Box Styling */
    .stAlert {
        border-radius: 10px;
        padding: 15px;
        margin-bottom: 15px;
    }
    .stInfo {
        background-color: #e0f2f7; /* Light blue info */
        border-left: 5px solid #2196F3;
        border-radius: 8px;
        padding: 15px;
    }
    .stSuccess {
        background-color: #e8f5e9; /* Light green success */
        border-left: 5px solid #4CAF50;
        border-radius: 8px;
        padding: 15px;
    }
    .stWarning {
        background-color: #fff3e0; /* Light orange warning */
        border-left: 5px solid #FF9800;
        border-radius: 8px;
        padding: 15px;
    }

    /* Metric styling based on value */
    /* Note: Streamlit's internal classes for metrics can be unstable,
             but these are common targets. */
    .st-emotion-cache-1f03409.e1nzilvr1 div[data-testid="stMetricValue"] {
        font-size: 2.8em; /* Larger value */
        font-weight: bold;
        color: #4CAF50; /* Default green for score */
        transition: color 0.5s ease; /* Smooth color transition */
    }
    .st-emotion-cache-1f03409.e1nzilvr1 div[data-testid="stMetricLabel"] {
        font-size: 1.3em; /* Larger label */
        color: #555;
    }
    .st-emotion-cache-1f03409.e1nzilvr1 div[data-testid="stMetricDelta"] {
        font-size: 1.6em; /* Larger percentage */
        color: #2196F3; /* Blue for percentage */
    }

    /* Matched keywords styling */
    .matched-keywords-box {
        background-color: #e0f7fa; /* Light cyan */
        border-left: 4px solid #00BCD4; /* Cyan border */
        padding: 15px; /* More padding */
        border-radius: 8px;
        margin-top: 15px; /* More margin */
        word-wrap: break-word;
        box-shadow: 0 2px 4px rgba(0,0,0,0.08); /* Subtle shadow */
        transition: all 0.3s ease; /* Smooth transition */
    }
    .matched-keywords-box:hover {
        transform: translateY(-2px); /* Slight lift on hover */
        box-shadow: 0 4px 8px rgba(0,0,0,0.12);
    }

    /* Keyframe Animations */
    @keyframes fadeIn {
        from { opacity: 0; }
        to { opacity: 1; }
    }
    @keyframes slideInLeft {
        from { transform: translateX(-100%); opacity: 0; }
        to { transform: translateX(0); opacity: 1; }
    }
    @keyframes pulse {
        0% { transform: scale(1); }
        50% { transform: scale(1.02); }
        100% { transform: scale(1); }
    }

    /* Apply pulse to results on load */
    .st-emotion-cache-f1g0g0.e1f1d6gn0 > div { /* This targets the outer div of each result block */
        animation: pulse 1.5s ease-in-out infinite;
    }

    /* Responsive adjustments */
    @media (max-width: 768px) {
        h1 { font-size: 2em; }
        h2 { font-size: 1.6em; }
        h3 { font-size: 1.3em; }
        .stButton>button { padding: 10px 20px; font-size: 15px; }
        .st-emotion-cache-1f03409.e1nzilvr1 div[data-testid="stMetricValue"] { font-size: 2em; }
        .st-emotion-cache-1f03409.e1nzilvr1 div[data-testid="stMetricLabel"] { font-size: 1em; }
        .st-emotion-cache-1f03409.e1nzilvr1 div[data-testid="stMetricDelta"] { font-size: 1.2em; }
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- 1. Text Extraction Functions ---

def extract_text_from_pdf(pdf_file):
    """Extracts text from a PDF file."""
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text += page.extract_text() or "" # Add empty string if page is empty
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(docx_file):
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(docx_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
    return text

def extract_text_from_txt(txt_file):
    """Extracts text from a TXT file."""
    return txt_file.read().decode("utf-8")

# --- 2. Text Processing and Keyword Matching ---

def preprocess_text(text):
    """Converts text to lowercase, removes non-alphanumeric characters, removes stopwords."""
    text = text.lower()
    # Keep only letters, numbers, and spaces
    text = re.sub(r'[^a-z0-9\s]', '', text)
    words = text.split()
    
    # Stopwords removal is optional now, uncomment if you want to use it
    # Make sure you have run nltk.download('stopwords') beforehand
    # try:
    #     stop_words = set(stopwords.words('english'))
    #     words = [word for word in words if word not in stop_words]
    # except LookupError:
    #     st.warning("NLTK stopwords data not found. Please run 'python -c \"import nltk; nltk.download(\'stopwords\')\"' in your terminal.")
    # except Exception as e:
    #     st.error(f"Error processing stopwords: {e}")

    return " ".join(words)

def calculate_keyword_match_score(resume_text, job_description_keywords):
    """
    Finds job description keywords in the resume text
    and calculates a match score.
    """
    score = 0
    matched_keywords = []

    # Ensure both keywords and resume text are processed
    processed_resume_text = preprocess_text(resume_text)
    # Ensure job_description_keywords are also preprocessed before checking
    processed_jd_keywords = [preprocess_text(kw) for kw in job_description_keywords if kw.strip()] # Remove empty keywords

    # Check for each required keyword if it exists in the resume
    for keyword in processed_jd_keywords:
        # Check for exact substring match using 'in' operator
        if keyword in processed_resume_text:
            score += 1
            matched_keywords.append(keyword)
    return score, matched_keywords

# --- Streamlit UI ---

st.set_page_config(layout="wide")
st.title("👨‍💻 Resume Screening App (Keyword Matching)")
st.markdown("---")

# Sidebar for inputs
st.sidebar.header("⚙️ Settings and Upload")

st.sidebar.subheader("1. Enter Job Requirements")
job_keywords_input = st.sidebar.text_area(
    "Enter required skills/keywords for the job, separated by commas or new lines:",
    "Python, Salesforce, Reactjs, Nodejs, Git, Machine Learning, Data Science, SQL, Cloud, Communication, Problem Solving"
)
# Split keywords into a list and remove empty strings
required_keywords = [kw.strip() for kw in job_keywords_input.split(',') if kw.strip()]
if not required_keywords:
    st.sidebar.warning("Please enter job keywords.")
else:
    st.sidebar.info(f"**Identified Keywords:** {', '.join(required_keywords)}")

st.sidebar.subheader("2. Upload Resumes")
uploaded_resumes = st.sidebar.file_uploader(
    "Upload resumes in PDF, DOCX or TXT format",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=True
)

# Main content area
st.header("📊 Screening Results")

if uploaded_resumes and required_keywords:
    st.markdown("---")
    results = []
    for uploaded_file in uploaded_resumes:
        file_name = uploaded_file.name
        file_type = file_name.split('.')[-1].lower()
        resume_text = ""

        with st.spinner(f"Extracting text from '{file_name}'..."):
            if file_type == "pdf":
                resume_text = extract_text_from_pdf(uploaded_file)
            elif file_type == "docx":
                resume_text = extract_text_from_docx(uploaded_file)
                # Reset stream position for DOCX file to allow re-reading if needed
                uploaded_file.seek(0)
            elif file_type == "txt":
                resume_text = extract_text_from_txt(uploaded_file)
            else:
                st.warning(f"Unsupported file type: {file_type} ({file_name})")
                continue

        if resume_text:
            score, matched_kws = calculate_keyword_match_score(resume_text, required_keywords)
            results.append({
                "file_name": file_name,
                "score": score,
                "matched_keywords": matched_kws,
                "total_required_keywords": len(required_keywords)
            })
        else:
            st.warning(f"Could not extract text from '{file_name}'.")

    # Sort results by score (highest score first)
    results.sort(key=lambda x: x["score"], reverse=True)

    if results:
        for i, res in enumerate(results):
            st.subheader(f"📄 {i+1}. {res['file_name']}")
            
            col1, col2 = st.columns(2) # Two columns for layout

            with col1:
                # Calculate percentage for delta
                percentage_match = round((res['score'] / res['total_required_keywords']) * 100, 2) if res['total_required_keywords'] > 0 else 0
                
                # Determine color for metric value based on percentage
                metric_color = ""
                if percentage_match >= 75:
                    metric_color = "#4CAF50" # Green
                elif percentage_match >= 50:
                    metric_color = "#FFC107" # Amber
                else:
                    metric_color = "#F44336" # Red

                st.markdown(
                    f"""
                    <div style="text-align: center;">
                        <div style="font-size: 1.2em; color: #555;">Match Score</div>
                        <div style="font-size: 2.5em; font-weight: bold; color: {metric_color};">{res['score']} / {res['total_required_keywords']}</div>
                        <div style="font-size: 1.5em; color: #2196F3;">{percentage_match}%</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            
            with col2:
                if res['matched_keywords']:
                    st.write(f"**Matched Keywords:**")
                    st.markdown(f"<div class='matched-keywords-box'>{', '.join(res['matched_keywords'])}</div>", unsafe_allow_html=True)
                else:
                    st.write("**No keywords matched.**")
            
            st.divider() # Visual separator
    else:
        st.info("No resumes uploaded or could not be processed.")
else:
    st.info("To get started, enter job keywords in the sidebar and upload resumes.")
